import os
import google.generativeai as genai
from dotenv import load_dotenv
from pathlib import Path
from docx import Document
import datetime
import re
import requests
from urllib.parse import urlparse
from typing import Optional
from bs4 import BeautifulSoup

# 環境変数を読み込み
env_path = Path(__file__).parent.parent.parent.parent / 'env.local'
load_dotenv(dotenv_path=env_path)
print(f"DEBUG (gemini.py): Loading env from {env_path}")

def summarize(text):
    """
    指定の思考ステップで最高の要約を出力する。
    """
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError(f"APIキーが見つかりません。{env_path} に GEMINI_API_KEY='YOUR_KEY' を設定してください。")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
        prompt = f"""
# 条件:
以下のステップを踏み、最終的な回答のみを {{出力フォーマット}} に従って出力してください。
中間プロセス（解釈・初期回答・再考・自己評価）は表示しないでください。

1. {{質問}} を内部で解釈してください。
2. 回答を考えてください。
3. 回答が適切か再考してください。
4. 最終的な回答を生成してください。
5. 自己評価は内部で行ってください。

# 出力フォーマット:
**回答**:
（最終的な回答のみを簡潔・明確に記述）

# 質問:
あなたは、プロの編集者です。最高の要約を出力してください。

# 要約対象:
{text}
"""
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"エラーが発生しました: {e}"


def generate_daily_report(text):
    """
    経営管理部長向け日報フォーマットで要約を生成する。
    """
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError(f"APIキーが見つかりません。{env_path} に GEMINI_API_KEY='YOUR_KEY' を設定してください。")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
        output_format_instructions = """
```
# 日報要約テンプレート（Markdown出力形式）

# 提出者情報
report_date: "2025-07-15"

# 1. ハイライト（重要トピック／対応事項の要約）
highlights:
  - "[具体的な業務名]：対応内容、関係者、状況（例：電子カルテB2接続障害の対応、天谷・吉田が処置）"
  - "[例：アプリ不具合対応]：機種・バージョン、原因、担当者、今後の対応予定"
  - "[例：対外対応]：業者名、日程、進捗状況"

# 2. 業務報告（当日実施した主な業務を分類別に記録）
tasks:
  - category: "電子カルテ・システム・IT"
    items:
      - "[業務名]：目的、実施内容、関係者、成果（例：スマートフォンOS更新）"
  - category: "総務・契約・調整"
    items:
      - "[業務名]：相手先、調整事項、進捗（例：システム導入契約の進捗共有）"
  - category: "その他院内対応"
    items:
      - "[業務名]：病棟・部門名、内容、担当、特記事項"

# 3. 継続検討・課題
ongoing_issues:
  - "[課題名]：現在の状況、阻害要因、次アクション（例：恒久対策の検討）"
  - "[例：設定変更未完了]：対象部署、原因、予定対応日"

# 4. 人材育成・指導対応（あれば記載）
people_management:
  - "[例：部下の対応支援]：対象者、場面、対応内容、所感"
  - "[例：指示内容]：目的、実施内容、反応"

# 5. その他連絡・情報共有
shared_info:
  - "[例：社内連絡]：担当者、内容（例：HP更新業務引継ぎ日程）"
  - "[例：会議日程]：タイトル、出席者、調整内容"
  - "[例：資料作成・更新]：ファイル名、保存先（OneDrive）、利用目的"

# 6. コメント・所感（自由記述欄）
comments: >
  [例：今日の全体的な所感や、改善点、気づきなど。不要な場合は記載なしでも可]

```
"""
        prompt = f"""
以下の指示に従い、日報を生成してください。

{output_format_instructions}

--- 入力データ ---
{text}
---
"""
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"エラーが発生しました: {e}"


def gemini_generate_content(prompt, model_name="gemini-2.0-flash"):
    """
    任意プロンプトでGemini APIからテキスト生成
    """
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError(f"APIキーが見つかりません。{env_path} に GEMINI_API_KEY='YOUR_KEY' を設定してください。")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"エラーが発生しました: {e}"


def create_word_report(sections, filename, save_dir):
    """
    セクションごとのテキストをWordファイルにまとめて保存
    sections: [(見出し, 本文), ...]
    """
    doc = Document()
    for title, content in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    save_path = os.path.join(save_dir, filename)
    doc.save(save_path)
    return save_path


def interactive_report_workflow() -> None:
    """
    Run an interactive workflow to gather research on a topic or URL and
    generate a Word document in a OneDrive‑synced folder.

    The user is prompted for a topic (either free text or a URL) and a
    purpose for the resulting report. If a URL is provided, the page
    content is fetched and analysed. The function then calls out to
    ``gemini_generate_content`` for research, structure creation,
    body generation and summary. Finally, it writes the sections to a
    Word document via ``create_word_report`` and saves it to a OneDrive
    directory.
    """

    def is_url(text: str) -> bool:
        """Return True if ``text`` looks like an HTTP or HTTPS URL."""
        url_pattern = re.compile(
            r'^(https?://[^\s]+)$'
        )
        return bool(url_pattern.match(text.strip()))

    def fetch_url_content(url: str) -> str:
        """Fetch the raw text of a URL with a sensible User‑Agent header.

        This helper attempts to retrieve the content of a remote HTTP
        resource. It sets a modern browser User‑Agent header, honours
        the declared encoding, and returns a string. Any network or
        HTTP errors are caught and returned as a human‑readable
        message.
        """
        try:
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/113.0.0.0 Safari/537.36"
                )
            }
            resp = requests.get(url, headers=headers, timeout=10)
            # Raise an exception for HTTP status codes >= 400
            resp.raise_for_status()
            # Attempt to detect the correct encoding
            resp.encoding = resp.apparent_encoding or resp.encoding
            return resp.text
        except Exception as e:
            return f"（URL取得エラー: {e}）"

    def sanitize_filename(filename: str) -> str:
        """
        Remove or replace characters that are unsafe for filenames.

        Only allow Japanese characters, alphanumerics and a subset of
        punctuation. Everything else becomes an underscore.
        """
        return re.sub(r'[^\w\u3000-\u303F\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF\-_.]', '_', filename)

    # Prompt the user for inputs
    print("=== 調査資料自動生成ワークフロー ===")
    topic = input("調査したいテーマやキーワード、またはURLを入力してください: ")
    purpose = input("資料の用途や目的を簡単に入力してください: ")

    url_mode = is_url(topic)
    url_content: Optional[str] = None
    url_for_filename: Optional[str] = None

    # Prepare prompts depending on whether a URL was provided
    if url_mode:
        print("URLが入力されました。Webサイト内容を取得しています...")
        url_content = fetch_url_content(topic)
        # Derive a safe filename fragment from the URL's host/path
        parsed = urlparse(topic)
        url_for_filename = parsed.netloc + parsed.path
        url_for_filename = url_for_filename.replace("/", "_").replace("\\", "_")
        if len(url_for_filename) > 50:
            url_for_filename = url_for_filename[:50]
        # Prepare prompts for Gemini
        research_prompt = (
            f"次のWebページの内容を解析し、{purpose}のための最新リサーチ結果を要点箇条書きでまとめてください。\n"
            f"--- Webページ内容 ---\n{url_content}\n---"
        )
        structure_prompt = (
            f"次のWebページの内容をもとに、{purpose}のための資料構成案（章立て・見出し）を作成してください。\n"
            f"--- Webページ内容 ---\n{url_content}\n---"
        )
        body_prompt = (
            f"次のWebページの内容をもとに、以下の構成案に沿って本文を執筆してください。\n"
            f"構成案: [後述]\n"
            f"--- Webページ内容 ---\n{url_content}\n---"
        )
        summary_prompt = (
            f"次のWebページの内容をもとに、要点を200字以内で要約してください。用途: {purpose}\n"
            f"--- Webページ内容 ---\n{url_content}\n---"
        )
    else:
        research_prompt = f"『{topic}』について、{purpose}のための最新リサーチ結果を要点箇条書きでまとめてください。"
        structure_prompt = f"『{topic}』について、{purpose}のための資料構成案（章立て・見出し）を作成してください。"
        body_prompt = f"『{topic}』について、以下の構成案に沿って本文を執筆してください。\n構成案: [後述]"
        summary_prompt = f"『{topic}』の要点を200字以内で要約してください。用途: {purpose}"

    # Invoke Gemini to generate each section
    print("Geminiでリサーチ中...")
    research = gemini_generate_content(research_prompt)
    print("構成案を生成中...")
    structure = gemini_generate_content(structure_prompt)
    print("本文を生成中...")
    # Replace the placeholder [後述] with the actual structure
    if "[後述]" in body_prompt:
        body_prompt = body_prompt.replace("[後述]", structure)
    body = gemini_generate_content(body_prompt)
    print("要約を生成中...")
    summary = gemini_generate_content(summary_prompt)

    def strip_markdown(text: str) -> str:
        """
        Remove common Markdown syntax from a string, returning plain text.

        This helper performs a best‑effort conversion by removing code
        fences, inline code markers, emphasis markers, headers, list
        markers and links/images. It does not attempt to interpret
        complex Markdown constructs; the goal is simply to avoid
        leaving raw Markdown markup in the Word document when no
        Markdown parser is available.
        """
        # Remove fenced code blocks
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
        # Remove inline code
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # Remove images: ![alt](url) -> alt
        text = re.sub(r'!\[(.*?)\]\([^)]*\)', r'\1', text)
        # Remove links: [text](url) -> text
        text = re.sub(r'\[(.*?)\]\([^)]*\)', r'\1', text)
        # Remove strong/bold/italic markers **text**, __text__, *text*, _text_
        text = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text)
        text = re.sub(r'(\*|_)(.*?)\1', r'\2', text)
        # Remove headings '#' at start of line
        text = re.sub(r'^\s*#{1,6}\s+', '', text, flags=re.MULTILINE)
        # Remove list markers '-', '*', '+', or numbered lists at line start
        text = re.sub(r'^\s*([\-*+]\s+|\d+\.\s+)', '', text, flags=re.MULTILINE)
        # Replace multiple consecutive blank lines with a single blank line
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    # Convert all sections to plain text to avoid leaving Markdown in Word
    research_plain = strip_markdown(research)
    structure_plain = strip_markdown(structure)
    body_plain = strip_markdown(body)
    summary_plain = strip_markdown(summary)

    # Determine a safe filename based on the topic or URL
    today = datetime.date.today().strftime("%Y%m%d")
    if url_mode:
        parsed_url = urlparse(topic)
        domain = parsed_url.netloc.replace('.', '_')
        # Attempt to analyse the page title or content if BeautifulSoup is available
        try:
            # Use the same headers as fetch_url_content for consistency
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/113.0.0.0 Safari/537.36"
                )
            }
            resp = requests.get(topic, headers=headers, timeout=10)
            resp.raise_for_status()
            resp.encoding = resp.apparent_encoding or resp.encoding
            page_title = ''
            main_content = ''
            if BeautifulSoup:
                soup = BeautifulSoup(resp.text, 'html.parser')
                # Title
                if soup.title and soup.title.string:
                    page_title = soup.title.string.strip()
                # Main text: prefer <main> tag, fall back to <body>
                if soup.main:
                    main_content = soup.main.get_text(separator='\n', strip=True)
                elif soup.body:
                    main_content = soup.body.get_text(separator='\n', strip=True)
                else:
                    main_content = soup.get_text(separator='\n', strip=True)
                main_content = main_content[:3000]
            # Choose a base name using Gemini if possible
            if main_content:
                prompt = (
                    "次のWebページ内容を分析し、資料ファイル名にふさわしい簡潔な日本語タイトルを20文字以内で生成してください。"
                    "記号やスラッシュ、コロンは使わず、内容が分かるようにしてください。\n"
                    f"内容: {main_content}"
                )
                analyzed_title = gemini_generate_content(prompt)
                analyzed_title = analyzed_title.split('\n')[0].strip()
                base_name = analyzed_title if analyzed_title else domain
            elif page_title:
                prompt = (
                    "次のWebページタイトルから、資料ファイル名にふさわしい簡潔な日本語タイトルを20文字以内で生成してください。"
                    "記号やスラッシュ、コロンは使わず、内容が分かるようにしてください。\n"
                    f"タイトル: {page_title}"
                )
                analyzed_title = gemini_generate_content(prompt)
                analyzed_title = analyzed_title.split('\n')[0].strip()
                base_name = analyzed_title if analyzed_title else domain
            else:
                base_name = domain
        except Exception:
            # Fall back to using the domain if we fail to fetch or parse
            base_name = domain
        safe_topic = sanitize_filename(base_name)
    else:
        safe_topic = sanitize_filename(topic)
    filename = f"{today}_{safe_topic}_資料.docx"

    # Define the OneDrive path and ensure it exists
    onedrive_dir = os.path.expanduser(
        "/Users/ueshima/Library/CloudStorage/OneDrive-医療法人社団　慶友会　吉田病院/00_Agents/資料"
    )
    os.makedirs(onedrive_dir, exist_ok=True)
    # Build the list of sections with plain text. Passing plain text
    # avoids including Markdown markup in the final Word document.
    sections = [
        ("リサーチ", research_plain),
        ("構成案", structure_plain),
        ("本文", body_plain),
        ("要約", summary_plain),
    ]
    save_path = create_word_report(sections, filename, onedrive_dir)
    print(f"\n資料が作成されました: {save_path}")
    print("OneDriveで自動同期されます。")


def short_video_workflow() -> None:
    """
    病院内レストランの収益アップのためのショート動画作成ワークフローを自動実行し、
    各工程ごとにGeminiでテキスト生成し、WordファイルにまとめてOneDriveに保存する。
    """
    import datetime
    print("=== 病院内レストラン ショート動画作成ワークフロー ===")
    topic = "病院内レストランの収益アップ施策"
    today = datetime.date.today().strftime("%Y%m%d")
    # 各工程のプロンプト
    prompts = [
        ("調査", f"『{topic}』についてリサーチし、要点をまとめてください。"),
        ("構成案", "ショート動画の構成案を作成してください。"),
        ("本文", "ショート動画用のナレーション原稿を作成してください。"),
        ("要約", "動画内容をSNS投稿用に要約してください。"),
        ("SNS投稿文", "Instagram用の投稿文を作成してください。"),
        ("画像生成プロンプト", "サムネイル画像生成用のプロンプトを作成してください。"),
        ("音声台本", "ナレーション用の台本を作成してください。"),
    ]
    sections = []
    for title, prompt in prompts:
        print(f"{title}を生成中...")
        content = gemini_generate_content(prompt)
        # Markdown除去
        content_plain = re.sub(r'[`*_\[\]#\-]', '', content)
        sections.append((title, content_plain))
    # 保存先
    safe_topic = sanitize_filename(topic)
    filename = f"{today}_{safe_topic}_ショート動画案.docx"
    onedrive_dir = os.path.expanduser(
        "/Users/ueshima/Library/CloudStorage/OneDrive-医療法人社団　慶友会　吉田病院/00_Agents/資料"
    )
    os.makedirs(onedrive_dir, exist_ok=True)
    save_path = create_word_report(sections, filename, onedrive_dir)
    print(f"\nショート動画案が作成されました: {save_path}")
    print("OneDriveで自動同期されます。")